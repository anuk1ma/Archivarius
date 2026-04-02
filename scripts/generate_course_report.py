from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\Anuk1ma\Documents\HTML\BookProject")
SOURCE = ROOT / "docs" / "course-report-v1.md"
OUTPUT = ROOT / "docs" / "Курсовой_отчет_Archivarius_v1.docx"


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_text_paragraph(document, text="", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, style=None, first_line=True):
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    set_paragraph_format(paragraph, first_line=first_line)
    return paragraph


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=12)


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "").strip()


def parse_sections(lines):
    toc = []
    body = []
    current = None
    for line in lines:
        if line.startswith("## "):
            current = line[3:].strip()
        if current == "СОДЕРЖАНИЕ":
            if line.startswith("## ") and line[3:].strip() != "СОДЕРЖАНИЕ":
                current = "BODY"
            elif line.strip() and not line.startswith("## "):
                toc.append(clean_inline(line))
                continue
        if current == "BODY" or current not in (None, "СОДЕРЖАНИЕ"):
            body.append(line.rstrip("\n"))
    # remove everything before first body heading
    first_body_index = next((i for i, item in enumerate(body) if item.startswith("## ")), 0)
    return toc, body[first_body_index:]


def add_markdown_table(document, rows):
    parsed = []
    for row in rows:
        if re.match(r"^\|[-\s|]+\|$", row.strip()):
            continue
        cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    table = document.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.style = "Table Grid"
    for r_index, row in enumerate(parsed):
        for c_index, cell in enumerate(row):
            paragraph = table.cell(r_index, c_index).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(cell)
            set_run_font(run, size=12, bold=(r_index == 0))
            paragraph.paragraph_format.line_spacing = 1.0


def build_document():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    toc_items, body_lines = parse_sections(lines)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    add_text_paragraph(document, "КОЛЛЕДЖ ТОО «ASTANA IT UNIVERSITY»", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "Цикловая комиссия «Специальных дисциплин»", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "КУРСОВОЙ ПРОЕКТ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "06130100 Программное обеспечение (по видам)", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "4S06130103 Разработчик программного обеспечения", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(
        document,
        "по дисциплине «ПМ04 - Проектирование и обеспечение бесперебойной работы web-сайта»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(
        document,
        "на тему «Разработка системы поиска и заказа книг в библиотеке»",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "Выполнил: Марат Мансұр", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text_paragraph(document, "Студент группы: ПО2401", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text_paragraph(document, "Преподаватель: Сайлаубаева Наргиза Ержанқызы", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text_paragraph(document, "Астана 2026 г.", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    document.add_page_break()
    add_text_paragraph(document, "СОДЕРЖАНИЕ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    for item in toc_items:
        add_text_paragraph(document, item, first_line=False)

    document.add_page_break()

    table_buffer = []
    for line in body_lines:
        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue

        if table_buffer:
            add_markdown_table(document, table_buffer)
            table_buffer = []
            add_text_paragraph(document, "", first_line=False)

        clean = clean_inline(line)
        if not clean:
            add_text_paragraph(document, "", first_line=False)
            continue

        if clean.startswith("## "):
            add_text_paragraph(document, clean[3:], bold=True, style="Heading 1", first_line=False)
        elif clean.startswith("### "):
            add_text_paragraph(document, clean[4:], bold=True, style="Heading 2", first_line=False)
        elif re.match(r"^\d+\.\s", clean):
            add_text_paragraph(document, clean)
        elif clean.startswith("- "):
            add_text_paragraph(document, f"• {clean[2:]}")
        else:
            add_text_paragraph(document, clean)

    if table_buffer:
        add_markdown_table(document, table_buffer)

    document.save(str(OUTPUT))


if __name__ == "__main__":
    build_document()
    print("Saved report DOCX.")
