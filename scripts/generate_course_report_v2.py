from pathlib import Path
import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(r"C:\Users\Anuk1ma\Documents\HTML\BookProject")
SOURCE = ROOT / "docs" / "course-report-v2.md"
OUTPUT = ROOT / "docs" / "Курсовой_отчет_Archivarius_v3.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagrams" / "generated"

TITLE_LINES = [
    ("Министерство просвещения Республики Казахстан", WD_ALIGN_PARAGRAPH.CENTER),
    ("Колледж ТОО «Astana IT University»", WD_ALIGN_PARAGRAPH.CENTER),
    ("Цикловая комиссия «Специальных дисциплин»", WD_ALIGN_PARAGRAPH.CENTER),
]

TOC_ITEMS = [
    "ВВЕДЕНИЕ",
    "1. ЦЕЛЬ И ЗАДАЧИ ПРОЕКТА",
    "1.1 Цель проекта",
    "1.2 Задачи проекта",
    "2. ФУНКЦИОНАЛЬНЫЕ ВОЗМОЖНОСТИ ПРИЛОЖЕНИЯ",
    "3. ИСПОЛЬЗОВАННЫЕ СРЕДСТВА И ТЕХНОЛОГИИ",
    "4. ОПИСАНИЕ ИНТЕРФЕЙСОВ",
    "5. АРХИТЕКТУРНАЯ МОДЕЛЬ СИСТЕМЫ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
    "ПРИЛОЖЕНИЕ",
]

DIAGRAM_MAP = {
    "Рисунок 5.1": ("navigation.png", "Рисунок 5.1 - Структура и навигация web-приложения Archivarius"),
    "Рисунок 5.2": ("architecture.png", "Рисунок 5.2 - Архитектура web-приложения Archivarius"),
    "Рисунок 5.3": ("erd.png", "Рисунок 5.3 - ER-диаграмма базы данных системы Archivarius"),
    "Рисунок 5.4": ("use-case.png", "Рисунок 5.4 - Диаграмма вариантов использования системы Archivarius"),
    "Рисунок 5.5": ("order-flow.png", "Рисунок 5.5 - Схема процесса оформления заказа"),
}


def set_cell_border(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "10")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "8C6D52")


def set_run_font(run, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_body_paragraph(paragraph, first_line=True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Cm(1.25)


def add_paragraph(document, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=14, first_line=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    set_body_paragraph(paragraph, first_line=first_line)
    return paragraph


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=12)


def insert_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, fld_end])
    set_run_font(run, size=14)


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "").strip()


def parse_body(lines):
    in_body = False
    body = []
    for line in lines:
        if line.startswith("## ВВЕДЕНИЕ"):
            in_body = True
        if in_body:
            body.append(line.rstrip("\n"))
    return body


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    centered = text in {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", "ПРИЛОЖЕНИЕ", "СОДЕРЖАНИЕ"}
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.style = document.styles["Heading 1" if level == 1 else "Heading 2"]
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 14, bold=True)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def update_fields_with_word(doc_path: Path):
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(str(doc_path))
    try:
        for field in doc.Fields:
            try:
                field.Update()
            except Exception:
                pass
        for toc in doc.TablesOfContents:
            try:
                toc.Update()
            except Exception:
                pass
        doc.Save()
    finally:
        doc.Close()
        word.Quit()
    return True


def add_table(document, rows):
    parsed = []
    for row in rows:
        if re.match(r"^\|[-\s|]+\|$", row.strip()):
            continue
        parsed.append([cell.strip() for cell in row.strip().strip("|").split("|")])
    if not parsed:
        return
    table = document.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Mm(15), Mm(70), Mm(75), Mm(30)]
    for row_index, row in enumerate(parsed):
        for col_index, cell_text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col_index < len(widths):
                cell.width = widths[col_index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(cell_text)
            set_run_font(run, size=12, bold=(row_index == 0))


def add_placeholder(document, title, description):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(f"{title}\n{description}")
    set_run_font(run, size=13, bold=False)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_h = OxmlElement("w:tcW")
    tc_h.set(qn("w:w"), "9000")
    tc_h.set(qn("w:type"), "dxa")
    tc_pr.append(tc_h)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), "2200")
    tr_h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_h)


def add_diagram(document, marker):
    file_name, caption = DIAGRAM_MAP[marker]
    path = DIAGRAM_DIR / file_name
    if path.exists():
        try:
            stream = BytesIO(path.read_bytes())
            document.add_picture(stream, width=Cm(15.5))
            paragraph = document.paragraphs[-1]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_par = document.add_paragraph()
            caption_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_par.add_run(caption)
            set_run_font(run, size=12)
            caption_par.paragraph_format.line_spacing = 1.0
        except Exception:
            add_placeholder(document, caption, "Требуется вставить диаграмму.")
    else:
        add_placeholder(document, caption, "Требуется вставить диаграмму.")


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    body_lines = parse_body(lines)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    for text, align in TITLE_LINES:
        add_paragraph(document, text, align=align, first_line=False)

    for _ in range(4):
        add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_paragraph(document, "КУРСОВОЙ ПРОЕКТ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_paragraph(document, "06130100 - Программное обеспечение (по видам)", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_paragraph(document, "по дисциплине ПМ04 - Проектирование и обеспечение бесперебойной работы web-сайта", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_paragraph(document, "Тема: Разработка системы поиска и заказа книг в библиотеке", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    for _ in range(5):
        add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_paragraph(document, "Выполнил: Марат Мансұр", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    add_paragraph(document, "Студент группы ПО2401", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    add_paragraph(document, "Руководитель: Сайлаубаева Наргиза Ержанқызы", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)

    for _ in range(6):
        add_paragraph(document, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_paragraph(document, "Астана, 2026 год", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    document.add_page_break()
    add_heading(document, "СОДЕРЖАНИЕ", 1)
    toc_par = document.add_paragraph()
    toc_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_par.paragraph_format.line_spacing = 1.5
    insert_toc(toc_par)

    document.add_page_break()

    table_buffer = []
    for line in body_lines:
        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue

        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []
            add_paragraph(document, "", first_line=False)

        clean = clean_inline(line)
        if not clean:
            add_paragraph(document, "", first_line=False)
            continue

        if clean.startswith("## "):
            add_heading(document, clean[3:], 1)
            continue

        if clean.startswith("### "):
            add_heading(document, clean[4:], 2)
            continue

        if clean.startswith("[Место для диаграммы:") or clean.startswith("[Вставить диаграмму"):
            matched = None
            for marker in DIAGRAM_MAP:
                if marker in clean:
                    matched = marker
                    break
            if matched:
                add_diagram(document, matched)
            else:
                add_placeholder(document, "Место для диаграммы", clean)
            continue

        if clean.startswith("[Место для скриншота:"):
            parts = clean.strip("[]")
            title, _, description = parts.partition(" — ")
            add_placeholder(document, title.replace("Место для скриншота: ", ""), description)
            continue

        if re.match(r"^\d+\.\s", clean):
            add_paragraph(document, clean)
            continue

        if clean.startswith("- "):
            add_paragraph(document, f"• {clean[2:]}")
            continue

        add_paragraph(document, clean)

    if table_buffer:
        add_table(document, table_buffer)

    document.save(str(OUTPUT))
    update_fields_with_word(OUTPUT)
    print("Saved report DOCX.")


if __name__ == "__main__":
    build()
